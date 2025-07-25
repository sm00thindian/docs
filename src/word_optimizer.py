# src/word_optimizer.py

from .base_optimizer import BaseDocumentOptimizer
from docx import Document
from typing import Any, List

class WordDocumentOptimizer(BaseDocumentOptimizer):
    """
    Optimizer for Microsoft Word (.docx) documents.
    """
    def __init__(self, config_path: str = None):
        super().__init__(config_path)
        doc_config = self.config['document_types']['word']
        self.file_extension = doc_config['extension']
        self.chunk_size = doc_config['chunk_size']
        self.chunk_overlap = doc_config['chunk_overlap']

    def load_document(self, file_path: str) -> Document:
        return Document(file_path)

    def _extract_text(self, content: Document) -> str:
        # This method is overridden in optimize, but kept for compatibility
        return ''

    def optimize(self, content: Document) -> List[str]:
        """
        Optimizes the document: removes headers/footers, extracts paragraphs and tables as chunks.
        """
        chunks = []

        # Extract paragraphs from body (excluding headers/footers)
        for paragraph in content.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        # Extract tables as markdown chunks
        for table in content.tables:
            table_chunk = self._table_to_markdown(table)
            if table_chunk:
                chunks.append(table_chunk)

        # Optional: Clean each chunk
        chunks = [self._clean_text(chunk) for chunk in chunks]

        return chunks

    def _table_to_markdown(self, table) -> str:
        """
        Converts a table to markdown format.
        """
        markdown = []
        # Header row
        header = [cell.text.strip() for cell in table.rows[0].cells]
        markdown.append('| ' + ' | '.join(header) + ' |')
        markdown.append('| ' + '--- | ' * len(header))

        # Data rows
        for row in table.rows[1:]:
            data = [cell.text.strip() for cell in row.cells]
            markdown.append('| ' + ' | '.join(data) + ' |')

        return '\n'.join(markdown)

    def _clean_text(self, text: str) -> str:
        text = super()._clean_text(text)
        lines = text.split('\n')
        cleaned_lines = [line for line in lines if line.strip() and len(line.strip()) > 5]
        return '\n'.join(cleaned_lines)
